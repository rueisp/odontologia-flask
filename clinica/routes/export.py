# clinica/routes/export.py

# --- Importaciones Necesarias ---
from flask import Blueprint, send_file, request, render_template, current_app
from werkzeug.utils import secure_filename
from io import BytesIO
import pandas as pd
import os
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import pytz 
from sqlalchemy.orm import load_only  # <--- ESTA LÍNEA FALTABA

# --- Importaciones de tus Modelos ---
from ..models import db, Paciente, Evolucion

# --- Creación del Blueprint ---
export_bp = Blueprint('export', __name__)


# --- Exportar a Excel (OPTIMIZADO) ---
@export_bp.route('/exportar_excel/<int:id>')
def exportar_excel(id):
    # Cargar solo el paciente con los campos que vamos a exportar
    paciente = db.session.query(Paciente).options(
        load_only(
            Paciente.id,
            Paciente.primer_nombre,
            Paciente.segundo_nombre,
            Paciente.primer_apellido,
            Paciente.segundo_apellido,
            Paciente.tipo_documento,
            Paciente.documento,
            Paciente.fecha_nacimiento,
            Paciente.edad,
            Paciente.email,
            Paciente.telefono,
            Paciente.direccion,
            Paciente.barrio,
            Paciente.alergias,
            Paciente.motivo_consulta,
            Paciente.enfermedad_actual,
            Paciente.observaciones,
            Paciente.dentigrama_canvas,
            Paciente.imagen_perfil_url
        )
    ).get_or_404(id)
    
    datos = {"Campo": [], "Valor": []}
    
    # SOLO los campos que realmente existen y se usan
    campos_exportar = [
        ("ID", paciente.id),
        ("Primer Nombre", paciente.primer_nombre),
        ("Segundo Nombre", paciente.segundo_nombre),
        ("Primer Apellido", paciente.primer_apellido),
        ("Segundo Apellido", paciente.segundo_apellido),
        ("Tipo Documento", paciente.tipo_documento),
        ("Documento", paciente.documento),
        ("Fecha Nacimiento", paciente.fecha_nacimiento.strftime('%d/%m/%Y') if paciente.fecha_nacimiento else 'N/A'),
        ("Edad", paciente.edad),
        ("Email", paciente.email or 'N/A'),
        ("Teléfono", paciente.telefono),
        ("Dirección", paciente.direccion or 'N/A'),
        ("Barrio", paciente.barrio or 'N/A'),
        ("Alergias", paciente.alergias or 'No especificado'),
        ("Motivo Consulta", paciente.motivo_consulta or 'No especificado'),
        ("Enfermedad Actual", paciente.enfermedad_actual or 'No especificado'),
        ("Observaciones", paciente.observaciones or 'No especificado')
    ]
    
    for campo, valor in campos_exportar:
        datos["Campo"].append(campo)
        datos["Valor"].append(valor if valor is not None else "No disponible")
        
    df = pd.DataFrame(datos)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Paciente')
    output.seek(0)
    
    filename = f"Paciente_{paciente.documento or paciente.id}.xlsx"
    return send_file(output, download_name=filename, as_attachment=True)


# --- Exportar a Word (OPTIMIZADO) ---
@export_bp.route('/exportar_word/<int:id>')
def exportar_word(id):
    # Cargar paciente con campos necesarios
    paciente = db.session.query(Paciente).options(
        load_only(
            Paciente.id,
            Paciente.primer_nombre,
            Paciente.segundo_nombre,
            Paciente.primer_apellido,
            Paciente.segundo_apellido,
            Paciente.tipo_documento,
            Paciente.documento,
            Paciente.fecha_nacimiento,
            Paciente.edad,
            Paciente.email,
            Paciente.telefono,
            Paciente.direccion,
            Paciente.barrio,
            Paciente.alergias,
            Paciente.motivo_consulta,
            Paciente.enfermedad_actual,
            Paciente.observaciones
        )
    ).get_or_404(id)
    
    # Cargar evoluciones del paciente
    evoluciones = db.session.query(Evolucion).options(
        load_only(Evolucion.fecha, Evolucion.descripcion)
    ).filter(Evolucion.paciente_id == id).order_by(Evolucion.fecha.asc()).all()
    
    doc = Document()

    # Estilo de fuente por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8) 

    # Zona Horaria
    local_timezone = pytz.timezone('America/Bogota') 
    now_in_local_tz = datetime.now(local_timezone)

    # --- ENCABEZADO ---
    fecha_emision_formateada = now_in_local_tz.strftime('%d/%m/%Y %H:%M')
    p_fecha = doc.add_paragraph(f'Fecha de Emisión: {fecha_emision_formateada}')
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_fecha.runs:
        run.font.size = Pt(7)
        run.italic = True

    titulo = doc.add_heading('Historia Clínica Odontológica', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_consultorio = subtitulo.add_run('Odontologia Dr. Rueis Pitre')
    font_consultorio = run_consultorio.font
    font_consultorio.name = 'Calibri'
    font_consultorio.size = Pt(10)
    font_consultorio.italic = True
    font_consultorio.bold = True

    doc.add_paragraph() 

    # --- DATOS DE FILIACIÓN (SOLO CAMPOS ESENCIALES) ---
    doc.add_heading('1. Datos de Filiación', level=2)
    campos_filiacion = [
        ("Nombres", f"{paciente.primer_nombre} {paciente.segundo_nombre or ''}".strip()),
        ("Apellidos", f"{paciente.primer_apellido} {paciente.segundo_apellido or ''}".strip()),
        ("Tipo Doc.", paciente.tipo_documento),
        ("Documento", paciente.documento),
        ("Fecha Nac.", paciente.fecha_nacimiento.strftime('%d/%m/%Y') if paciente.fecha_nacimiento else 'N/A'),
        ("Edad", paciente.edad),
        ("Email", paciente.email or 'N/A'),
        ("Teléfono", paciente.telefono),
        ("Dirección", paciente.direccion or 'N/A'),
        ("Barrio", paciente.barrio or 'N/A')
    ]
    crear_tabla_formato(doc, campos_filiacion, una_columna=False, label_font_size=Pt(7), value_font_size=Pt(7), vertical_align_top=True)
    doc.add_paragraph()

    # --- INFORMACIÓN CLÍNICA ---
    doc.add_heading('2. Información Clínica', level=2)
    campos_clinicos = [
        ("Motivo de Consulta", limpiar_texto_para_word(paciente.motivo_consulta)),
        ("Enfermedad Actual", limpiar_texto_para_word(paciente.enfermedad_actual)),
        ("Alergias", limpiar_texto_para_word(paciente.alergias)),
        ("Observaciones", limpiar_texto_para_word(paciente.observaciones))
    ]
    crear_tabla_formato(doc, campos_clinicos, una_columna=False, label_font_size=Pt(7), value_font_size=Pt(7), vertical_align_top=True)
    doc.add_paragraph()

    # --- EVOLUCIÓN ---
    doc.add_heading('3. Evolución del Paciente', level=2)
    if evoluciones:
        tabla_evos = doc.add_table(rows=1, cols=2)
        tabla_evos.style = 'Table Grid'
        tabla_evos.columns[0].width = Inches(1.25)
        tabla_evos.columns[1].width = Inches(5.25)
        hdr_cells = tabla_evos.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].text = 'Descripción de la Evolución'
        hdr_cells[1].paragraphs[0].runs[0].bold = True

        for evo in evoluciones:
            row_cells = tabla_evos.add_row().cells
            
            # Manejo de zona horaria
            if evo.fecha.tzinfo is None: 
                fecha_evo_utc = evo.fecha.replace(tzinfo=pytz.utc)
            else: 
                fecha_evo_utc = evo.fecha

            fecha_evo_local = fecha_evo_utc.astimezone(local_timezone)
            row_cells[0].text = fecha_evo_local.strftime('%d/%m/%Y %H:%M') 
            row_cells[1].text = evo.descripcion
    else:
        doc.add_paragraph("No hay evoluciones registradas.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    download_filename = f"Historia_Clinica_{paciente.documento or paciente.id}.docx"
    return send_file(output, download_name=download_filename, as_attachment=True)


# --- FUNCIONES AUXILIARES ---
def limpiar_texto_para_word(texto):
    if texto is None: return ""
    texto = str(texto)
    texto = texto.replace('\r\n', ' ').replace('\n', ' ')
    texto = re.sub(r'\s+', ' ', texto)
    return texto.strip()

def crear_tabla_formato(doc, campos, una_columna=False, label_font_size=None, value_font_size=None, vertical_align_top=False):
    cols = 2 if una_columna else 4
    tabla = doc.add_table(rows=0, cols=cols)
    tabla.style = 'Table Grid'

    if una_columna:
        tabla.columns[0].width = Inches(1.8)
        tabla.columns[1].width = Inches(4.7)
    else: 
        tabla.columns[0].width = Inches(1.2)
        tabla.columns[1].width = Inches(2.1)
        tabla.columns[2].width = Inches(1.2)
        tabla.columns[3].width = Inches(2.1)

    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    paso = 1 if una_columna else 2
    for i in range(0, len(campos), paso):
        row_cells = tabla.add_row().cells
        if vertical_align_top:
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        label_izq, value_izq = campos[i]
        p_label_izq = row_cells[0].paragraphs[0]
        run_label_izq = p_label_izq.add_run(f"{label_izq}:")
        run_label_izq.bold = True
        if label_font_size: run_label_izq.font.size = label_font_size

        p_value_izq = row_cells[1].paragraphs[0]
        run_value_izq = p_value_izq.add_run(str(value_izq) if value_izq is not None else 'N/A')
        if value_font_size: run_value_izq.font.size = value_font_size
        p_value_izq.paragraph_format.space_before = Pt(0)
        p_value_izq.paragraph_format.space_after = Pt(0)
        p_value_izq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_value_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if not una_columna and i + 1 < len(campos):
            label_der, value_der = campos[i + 1]
            p_label_der = row_cells[2].paragraphs[0]
            run_label_der = p_label_der.add_run(f"{label_der}:")
            run_label_der.bold = True
            if label_font_size: run_label_der.font.size = label_font_size

            p_value_der = row_cells[3].paragraphs[0]
            run_value_der = p_value_der.add_run(str(value_der) if value_der is not None else 'N/A')
            if value_font_size: run_value_der.font.size = value_font_size
            p_value_der.paragraph_format.space_before = Pt(0)
            p_value_der.paragraph_format.space_after = Pt(0)
            p_value_der.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_value_der.alignment = WD_ALIGN_PARAGRAPH.LEFT